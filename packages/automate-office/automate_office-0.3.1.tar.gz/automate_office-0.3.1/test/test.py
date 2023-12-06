from automate_office.automate_word.MyWordCreator import WordCreator
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION_START, WD_ORIENT

from automate_office.automate_word.helper_func import display_paragraphs, add_font_style, add_paragraph_style
from automate_office.automate_word.settings import DOC_FONT_SETTINGS

if __name__ == "__main__":
    my_word = WordCreator()

    p_s_1 = add_paragraph_style(my_word.doc, style_name="1", font_info=("微软雅黑", 12, False, False, False, (0, 0, 0)), first_line_indent=-1, line_spacing=0.5)
    p_s_2 = add_paragraph_style(my_word.doc, style_name="2", font_info=("微软雅黑", 16, False, False, False, (0, 0, 0)), first_line_indent=-4, line_spacing=4)


    my_word.add_paragraph(
        paragraph_info=[
            ["a1", None],
            ["a2", ("微软雅黑", 12, False, True, False, (255, 0, 0))],
            ["a3", None],
            ["a4", ("微软雅黑", 12, True, True, False, (0, 0, 255))],
            ["a5", None],
            ["a6" * 100, None],
        ],
        paragraph_style=p_s_1
    )

    my_word.add_paragraph(
        paragraph_info=[
            ["a1", ("微软雅黑", 12, True, False, False, (0, 0, 0))],
            ["a2", None],
            ["a3", ("微软雅黑", 12, False, False, True, (0, 255, 0))],
            ["a4", None],
            ["a5", ("微软雅黑", 12, False, True, True, (255, 255, 0))],
            ["a6" * 100, None],
        ],
        paragraph_style=p_s_2
    )
    my_word.save_doc()



