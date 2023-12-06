"""
    未完成，需要时再处理，无非是字体、段落、页眉页脚、表格、列表等的设置
"""

import pathlib

from docx.oxml.ns import qn
from docx import Document, section
from docx.shared import Inches, Cm
from docx.enum.section import WD_SECTION_START, WD_ORIENT

from automate_office.automate_word.helper_func import add_paragraph


class WordCreator(object):

    def __init__(self, section_config={}):
        self._create_doc(section_config=section_config)

    def _create_doc(self, section_config):
        self.doc =  Document()
        if section_config:
            self.set_section(cur_section=self.doc.sections[0], cur_section_config=section_config)
            
    
    def save_doc(self, filename=None, save_path=None):
        if save_path is None:
            save_path = pathlib.Path.home().joinpath("Downloads")
        elif not isinstance(save_path, pathlib.WindowsPath):
            try:
                save_path = pathlib.Path(save_path)
            except Exception as e:
                raise ValueError(f"存储路径错误，{e}")

            if not save_path.is_dir():
                raise ValueError(f"{save_path}不是目录")

        if filename is None:
            filename = "output.docx"

        self.doc.save(save_path.joinpath(filename))

    def set_section(
        self, 
        cur_section: section.Section, 
        cur_section_config={
            "unit_type": "inches",
            "params": {
                "top_margin": 1,
                "bottom_margin": 1,
                "left_margin": 1.25,
                "right_margin": 1.25,
                "page_height": 11,
                "page_width": 8.5,
            },
            # 目前不管用
            "orientation": WD_ORIENT.PORTRAIT
        }
    ):
        support_params = ("top_margin", "bottom_margin", "left_margin", "right_margin", "page_height", "page_width")

        unit_type = {
            "inches": Inches,
            "cm": Cm
        }.get(
            cur_section_config.get("unit_type", "inches"), Inches
        )
        
        setattr(cur_section, "orientation", cur_section_config.get("orientation", WD_ORIENT.PORTRAIT))

        for k, v in cur_section_config["params"].items():
            if k in support_params:
                setattr(cur_section, k, unit_type(v))

    def create_section(self, section_config={}, start_type=WD_SECTION_START.NEW_PAGE):
        new_section = self.doc.add_section(start_type=start_type)
        if section_config:
            self.set_section(cur_section=new_section, cur_section_config=section_config)
        return new_section

    @staticmethod
    def set_multiple_col(cur_section: section.Section, n_cols):
        """分栏: 最多支持45
        """
        cur_section._sectPr.xpath('./w:cols')[0].set(qn('w:num'), f'{n_cols:>.0f}') 
    
    def add_header(self, text, level=0):
        self.doc.add_heading(text, level)

    def add_paragraph(self, paragraph_info, paragraph_style=None):
  
        paragraph = self.doc.add_paragraph(style=paragraph_style)
        add_paragraph(cur_paragraph=paragraph, paragraph_info=paragraph_info)


    def 页眉页脚():
        from docx import Document
        document = Document()  # 新建文档  或 模板
        header = document.sections[0].header  # 获取第一个节的页眉
        print('页眉中默认段落数：', len(header.paragraphs))
        paragraph = header.paragraphs[0]  # 获取页眉的第一个段落
        paragraph.add_run('这是第一节的页眉2')  # 添加页面内容
        footer = document.sections[0].footer  # 获取第一个节的页脚
        paragraph = footer.paragraphs[0]  # 获取页脚的第一个段落
        paragraph.add_run('这是第一节的页脚2')  # 添加页脚内容
        document.save('test.docx')  # 保存文档

    def about_list(self):
        
        for i, s in enumerate(
            [
                "List", "List 2", "List 3", 
                "List Bullet", "List Bullet 2", "List Bullet 3", 
                "List Continue", "List Continue 2", "List Continue 3", 
                "List Number", "List Number 2", "List Number 3"
            ]
        ):
            _ = self.doc.add_paragraph(str(i) * 10, style=s)

        self.save_doc()