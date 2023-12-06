from docx.shared import Pt, Inches, RGBColor
from docx.enum.style import WD_STYLE_TYPE

from automate_office.automate_word.settings import DOC_FONT_SETTINGS


def set_font(cur_font, cur_font_info):
    """设置字体
    Args:
        cur_font (_type_): pptx.text.text.font类型
        cur_font_info (_type_): tuple, 字体、字号、加粗、斜体、颜色
    """
    font_name, font_size, bold, italic, underline, rgb = cur_font_info

    cur_font.name = font_name
    cur_font.size = Pt(font_size)
    cur_font.bold = bold
    cur_font.italic = italic
    cur_font.underline = underline
    cur_font.color.rgb = RGBColor(*rgb)


def add_font_style(document, style_name, font_name, font_size=12, font_color=(0, 0, 0), bold=False, italic=False, underline=False):
    # CHARACTER STYLE似乎现在不能生效
    styles = document.styles
    # 新建一个叫Citation的style
    style = styles.add_style(style_name, WD_STYLE_TYPE.CHARACTER)
    set_font(font=style.font, font_name=font_name, font_size=font_size, font_color=font_color, bold=bold, italic=italic, underline=underline)


def add_paragraph_style(document, style_name, font_info, first_line_indent, line_spacing):
    styles = document.styles
    style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    # 设置字体
    set_font(style.font, font_info)
    # 设置段落
    # paragraph_format.space_before = Pt(18)
    # paragraph_format.space_after = Pt(12)
    paragraph_format = style.paragraph_format
    paragraph_format.first_line_indent = Inches(first_line_indent)
    paragraph_format.line_spacing = Inches(line_spacing)

def add_paragraph(cur_paragraph, paragraph_info):

    for text, cur_font_info in paragraph_info:
        run = cur_paragraph.add_run(text)
        if cur_font_info is not None:
            set_font(run.font, cur_font_info)


def display_paragraphs(cur_doc):

    for i, p in enumerate(cur_doc.paragraphs):
        print(i, p.text)
