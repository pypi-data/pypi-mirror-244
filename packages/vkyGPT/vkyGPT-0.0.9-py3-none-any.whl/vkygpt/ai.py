from vkygpt import AsyncClient
from asyncio import run
import pyttsx3
import re
import psutil
from vkygpt.codedump import write_code
import speech_recognition as sr
import cv2
import ctypes
import subprocess
import psutil
from plyer import notification
from PIL import Image
import pytesseract
import youtube_dl
from googlesearch import *
import time
import tkinter
import random
import operator
import datetime
import webbrowser
import os
import pyttsx3
import winshell
import pyjokes
import feedparser
import requests
import shutil
import urllib.request
from bs4 import BeautifulSoup
from urllib.request import urlopen
from tkinter import filedialog
from tkinter import *
from pygame import mixer
import pyqrcode
from pyqrcode import QRCode
import math
import keyboard
from vkygpt.imagegen import imgen
# Initialize the text-to-speech engine
engine = pyttsx3.init()
voices = engine.getProperty('voices')
engine.setProperty('voice', voices[1].id)
engine.setProperty('rate', 190)

# Initialize the speech recognition
recognizer = sr.Recognizer()

def speak(audio):
	engine.say(audio)
	engine.runAndWait()

def clrs():
    clear = lambda: os.system('cls')
    clear()

def close_process(process_name):
    # Find the specified process and terminate it
    for process in psutil.process_iter(['pid', 'name']):
        if process_name in process.info['name'].lower():
            try:
                psutil.Process(process.info['pid']).terminate()
                print(f"{process_name.capitalize()} closed successfully.")
            except Exception as e:
                print(f"Error closing {process_name}: {e}")
            return

    print(f"{process_name.capitalize()} not found.")

async def process_code_blocks(prompt):
    print("🤖: Loading....")
    try:
        resp = await AsyncClient.create_completion("gpt3", prompt)
        # print(f"🤖: {resp}")
        programming_languages_lowercase = {
                'python': '.py',
                'javascript': '.js',
                'java': '.java',
                'c++': '.cpp',
                'c': '.c',
                'html': '.html',
                'css': '.css',
                'ruby': '.rb',
                'swift': '.swift',
                'php': '.php',
                'sql': '.sql',
                'typescript': '.ts',
                'shell scripting': '.sh',
                'go': '.go',
                'rust': '.rs',
                'kotlin': '.kt',
                'objective-c': '.m',
                'markdown': '.md',
                'perl': '.pl',
                'scala': '.scala',
                'dart': '.dart',
                'yaml': '.yaml',
                'xml': '.xml',
                'batch script': '.bat',
                'powershell': '.ps1',
                'r': '.r',
                'vue.js': '.vue',
                'jupyter notebooks': '.ipynb',
            }

        for lang, ext in programming_languages_lowercase.items():
            code_pattern = re.compile(fr'```{lang}(.*?)```', re.DOTALL | re.IGNORECASE)
            code_matches = code_pattern.findall(resp)
            for i, code_snippet in enumerate(code_matches, start=1):
                formatted_string = f"# Code snippet {i} in {lang}:\n{code_snippet.strip()}"
                write_code(formatted_string, prompt)
    except Exception as e:
        print(f"🤖: {e}")

async def main():
    method=input("How do you want to communicate ?(text/speech) : ")
    while True:
        try:
            if method == 'text'.lower() or method =='t'.lower():
                prompt= input("👦: Speak now...")
            else:
                with sr.Microphone() as source:
                    print("👦: Speak now...")
                    audio = recognizer.listen(source, timeout=5)
                    print("👦: Recognizing...")

                prompt = recognizer.recognize_google(audio)
                print(f"👦: {prompt}")
            a = prompt.lower()

            if 'quit' in a or 'bye' in a or 'bye Friday' in a:
                speak("Thank you Sir for your time. Have a nice day.")
                clrs()
                break

            elif 'hello' in a or 'hello Friday' in a or 'hi Friday' in a or 'hi' in a or 'hey Friday' in a or 'hey' in a:
                speak("Hello Sir. How are you ?")

            elif 'good morning' in a or 'good morning Friday' in a:
                speak("Good Morning Sir !")

            elif 'good evening' in a or 'good evening Friday' in a:
                speak("Good Evening Sir !")

            elif 'good night' in a or 'good night Friday' in a:
                speak("Good Night Sir !")

            elif 'open cmd' in a or 'open command prompt' in a or 'open cmd Friday' in a or 'open command prompt Friday' in a:
                speak("Opening command prompt")
                subprocess.run(["C:\\WINDOWS\\system32\\cmd.exe"], shell=True)

            elif 'how are you' in a or 'how do you do' in a or 'how are you Friday' in a or 'how do you do Friday' in a:
                speak("I am fine Sir. Thank you for asking.")

            elif 'what is your name' in a or 'who are you' in a or "what's your name" in a:
                speak("My name is Friday - The virtual assistant cum chatbot.")

            elif 'wait' in a or 'wait Friday' in a or 'wait please' in a or 'please wait' in a or 'please wait Friday' in a or 'wait please Friday' in a:
                speak("Ok Sir. As you wish ! Waiting till space is pressed !")
                keyboard.wait('space')

            elif 'what is the time' in a or 'what is time' in a or 'say the time' in a or 'say time' in a or 'show time' in a or 'open time' in a or "what's the time" in a:
                now = datetime.datetime.now()
                speak("Sir, the current time is "+now.strftime("%H:%M")+'\n')
                print("Sir, the current time is "+now.strftime("%H:%M")+'\n')

            elif 'get weather' in a or 'get the weather' in a or 'what is the weather' in a or 'give me the weather' in a or 'give me weather' in a:
                import requests, json

                # Enter your API key here
                api_key = "69bf0a590576448ed0bfd804ac2b2694"

                # base_url variable to store url
                base_url = "http://api.openweathermap.org/data/2.5/weather?"

                speak("Sir please give the city name")
                city_name = a()

                complete_url = base_url + "appid=" + api_key + "&q=" + city_name

                # get method of requests module
                # return response object
                response = requests.get(complete_url)

                # json method of response object
                # convert json format data into
                # python format data
                x = response.json()

                if x["cod"] != "404":

                    y = x["main"]

                    current_temperature = y["temp"]

                    current_pressure = y["pressure"]

                    current_humidiy = y["humidity"]

                    z = x["weather"]

                    weather_description = z[0]["description"]
                    celsius = current_temperature-273.15

                    speak("Sir, Temperature in " + city_name + "is" + str(celsius) + " degree celsius")
                    print("Temperature in Celsius = " + str(celsius) + " degree celsius")
                    print("Temperature in Farenheit = " + str(celsius*9/5+32) + " degree farenheit")

                    speak("The weather description is " + str(weather_description))
                    print("Atmospheric pressure = " + str(current_pressure) + " mb")
                    print("Humidity = " + str(current_humidiy) + " %")
                    print("Description = " + str(weather_description))
                    print("")

                else:
                    print(" City Not Found. Please try again.")

            elif 'open google' in a or 'show google' in a:
                speak("Opening google")
                webbrowser.open("www.google.com")

            elif 'what is the date' in a or 'what is date' in a or 'show date' in a or 'show the date' in a or 'say the date' in a:
                now = datetime.datetime.now()
                speak("Sir, the Current Date is : "+now.strftime("%d-%m-%Y")+'\n')
                print ("Sir, the Current Date is : ")
                print (now.strftime("%d-%m-%Y")+'\n')
                print("")

            elif 'open calendar' in a or 'show calendar' in a:
                speak("Opening Calendar")
                subprocess.run(['start', 'outlookcal:'], shell=True)
                print("Opened Calendar")
                print("")

            elif 'close calendar' in a or 'exit calendar' in a:
                speak("Closing Calendar")
                close_process('calendar')

            elif 'open whatsapp' in a or 'show whatsapp' in a: ##Only if you have whatsapp web.
                speak("Opening Whatsapp")
                webbrowser.open("https://web.whatsapp.com/")

            elif 'open wolframalpha' in a or 'show wolframalpha' in a:
                speak("Opening Wolframaplha")
                webbrowser.open("https://www.wolframalpha.com/")

            elif 'open youtube' in a or 'show youtube' in a:
                speak("Opening Youtube")
                webbrowser.open("www.youtube.com")

            elif 'open gmail' in a or 'show gmail' in a or 'open g mail' in a or 'show g mail' in a:
                speak("Opening Gmail")
                webbrowser.open("https://mail.google.com/mail/u/0/#inbox")

            elif 'open notepad' in a or 'show notepad' in a:
                speak("Opening Notepad")
                subprocess.run(["C:\\WINDOWS\\system32\\notepad.exe"], shell=True)

            elif 'close notepad' in a or 'exit notepad' in a:
                speak("Closing notepad")
                close_process('notepad')

            elif 'open bluej' in a or 'show bluej' in a:## If you have Bluej only
                speak("Opening BlueJ")
                subprocess.run(["C:\\Program Files\\BlueJ\\BlueJ.exe"], shell=True)

            elif 'open python all' in a or 'show python all' in a or 'open pythonall' in a or 'show pythonall' in a:
                speak("Opening Python All")
                subprocess.run(["C:\\Python All"], shell=True)

            elif 'open zoom' in a or 'show zoom' in a:## If you have xoom only
                speak("Opening Zoom")
                subprocess.run(["C:\\Users\\joels\\AppData\\Roaming\\Zoom\\bin\\Zoom.exe"], shell=True)

            elif 'open news' in a or 'show news' in a or 'get news' in a or 'get the news' in a or 'what is the news' in a or 'news' in a:
                from newspaper import Article

                #A new article from TOI
                url = "https://timesofindia.indiatimes.com/news"

                toi_article = Article(url, language="en") # en for English

                toi_article.download()

                toi_article.parse()

                toi_article.nlp()
                speak("Sir, extracting news from"+toi_article.title)
                print(toi_article.title)

                print("")
                speak("Sir, below is the news")
                print(toi_article.text)
                print("")

            elif 'say a joke' in a or 'say joke' in a or 'tell a joke' in a:
                joke = pyjokes.get_joke()
                speak(joke)
                print(joke)
                print("")

            elif 'play a music' in a or 'play some music' in a or 'play the music' in a or 'play music' in a:
                speak("Here you go with music")
                subprocess.run("spotify", shell=True)
            elif 'close spotify' in a or 'close music' in a or 'exit music' in a or 'exit spotify' in a:
                speak("Closing Spotify")
                close_process('spotify')

            elif 'open calculator' in a or 'show calculator' in a:
                speak("Opening calculator.")
                subprocess.run("calc", shell=True)
            
            elif 'exit calculator' in a or 'close calculator' in a:
                speak("Closing Calculator")
                close_process('calc')

            elif 'create an image' in a:
                speak("Creating an Image")
                await imgen(prompt)
            elif 'take a pic' in a or 'take a picture' in a or 'open camera' in a or 'show camera' in a:
                speak("Preparing to take picture. Press s to save the image.")
                key = cv2. waitKey(1)
                webcam = cv2.VideoCapture(0)
                while True:
                    try:
                        check, frame = webcam.read()
                        print(check) #prints true as long as the webcam is running
                        print(frame) #prints matrix values of each framecd
                        cv2.imshow("Capturing", frame)
                        key = cv2.waitKey(1)

                        if key == ord('s'):
                            cv2.imwrite(filename='images\\saved_img.jpg', img=frame)
                            webcam.release()
                            img_new = cv2.imread('saved_img.jpg', cv2.IMREAD_GRAYSCALE)
                            img_new = cv2.imshow("Captured Image", img_new)
                            cv2.waitKey(1650)
                            cv2.destroyAllWindows()
                            print("Processing image...")
                            img_ = cv2.imread('saved_img.jpg', cv2.IMREAD_ANYCOLOR)
                            print("Converting RGB image to grayscale...")
                            gray = cv2.cvtColor(img_, cv2.COLOR_BGR2GRAY)
                            print("Converted RGB image to grayscale...")
                            print("Resizing image to 28x28 scale...")
                            img_ = cv2.resize(gray,(28,28))
                            print("Resized...")
                            img_resized = cv2.imwrite(filename='images\\saved_img-final.jpg', img=img_)
                            print("Image saved!")

                            break
                        elif key == ord('q'):
                            print("Turning off camera.")
                            webcam.release()
                            print("Camera off.")
                            print("Program ended.")
                            cv2.destroyAllWindows()
                            break

                    except(KeyboardInterrupt):
                        print("Turning off camera.")
                        webcam.release()
                        print("Camera off.")
                        print("Program ended.")
                        cv2.destroyAllWindows()
                        break

            elif 'take screenshot' in a or 'take a screenshot' in a or 'take the screenshot' in a:
                speak("Opening screenshot.")
                os.startfile("Screenshot Taker\\dist\\Screenshot Taker\\Screenshot Taker.exe")

            elif 'take a note' in a or 'take some note' in a or 'take note' in a or 'take the note' in a:
                from docx import Document
                from docx import Inches
                import time
                import datetime
                import os
                import tkinter
                from tkinter import filedialog

                root = tkinter.Tk()
                root.wm_withdraw()
                root.iconify()

                speak("Sir, what do you have to write on your notes ? Please give the input.")
                whattoright = a()

                now = datetime.datetime.now()

                documenttitle = now.strftime("Notes of %d-%m-%Y at %H:%M")
                document = Document()

                document.add_heading(documenttitle, 0)

                p = document.add_paragraph(whattoright).bold = True

                #document.add_picture('monty-truth.png', width=Inches(1.25))

                savingdocument = document.save(filedialog.asksaveasfilename(initialdir='C:\\Users\\joels\\Desktop\\Friday v2.0\\Documents'))

                os.startfile('Documents')

            elif 'lock the device' in a or 'lock device' in a or 'lock a device' in a or 'lock this device'  in a:
                speak("Locking the device")
                ctypes.windll.user32.LockWorkStation()
                break

            elif 'empty recycle bin' in a or 'empty bin' in a or 'empty trash' in a:
                winshell.recycle_bin().empty(confirm = False, show_progress = False, sound = True)
                speak("Trash items removed.")

            elif 'sent gmail' in a or 'sent the gmail' in a or 'sent g mail' in a or 'sent the g mail' in a or 'sent mail' in a:
                os.startfile("Senting Gmail\\dist\\Senting Gmail\\Senting Gmail.exe")

            elif 'Friday' in a or 'where are you Friday' in a:
                speak("At your service Sir !")

            elif 'open webex' in a or 'show webex' in a or 'open cisco webex' in a or 'show cisco webex' in a:
                speak("Opening Cisco Webex")
                webbrowser.open("https://globalpage-prod.webex.com/join?surl=https%3A%2F%2Fsignin.webex.com%2Fcollabs%2F%23%2Fmeetings%2Fjoinbynumber%3F")

            elif 'record Friday' in a or 'Friday record' in a or 'open recorder' in a or 'show recorder' in a:
                speak("Opening Voice Recorder")
                os.startfile("Recordingapp\\dist\\Recordingapp\\Recordingapp.exe")

            elif 'download video' in a or 'download a video' in a or 'download the video' in a:
                ydl_opts = {}

                def dwl_vid():
                    with youtube_dl.YoutubeDL(ydl_opts) as ydl:
                        ydl.download([zxt])

                channel = 1
                while (channel == int(1)):
                    speak("Copy & paste the URL of the YouTube video you want to download in the console")
                    link_of_the_video = input("Copy & paste the URL of the YouTube video you want to download:- ")
                    zxt = link_of_the_video.strip()

                    dwl_vid()
                    speak("Enter 1 if you want to download more videos or Enter 0 if you are done ")
                    channel = int(input("Enter 1 if you want to download more videos \nEnter 0 if you are done "))

            elif 'ocr' in a or 'extract text from image' in a:
                speak("Opening Optical Character Recognition OCR")
                root = tkinter.Tk()
                root.wm_withdraw()
                root.iconify()

                pytesseract.pytesseract.tesseract_cmd = "C:\\Users\\joels\\AppData\\Local\\Tesseract-OCR\\tesseract.exe"
                image = Image.open(filedialog.askopenfilename())

                image_to_text = pytesseract.image_to_string(image, lang='eng')
                print(image_to_text)

                root.destroy()

            elif 'qr code' in a or 'q r code' in a or 'qrcode' in a:
                root = tkinter.Tk()
                root.wm_withdraw()
                root.iconify()

                speak("What is the url for which you want the qr code ?")

                print("What is the url for which you want the qr code ?")
                s = input()
                url = pyqrcode.create(s)
                url.svg(filedialog.asksaveasfilename(initialdir = "Qr codes",title = "Save qr code",defaultextension='.svg'), scale = 8)

                speak("Qr code saved successfully !")
                root.destroy()

            elif 'show battery' in a or 'what is the battery remaining' in a or 'what is the remaining battery' in a:
                battery = psutil.sensors_battery()
                percent = battery.percent
                percentinletters = str(percent)
                speak("Sir, there is "+percentinletters+" percent remaining.")
                print("There is "+percentinletters+" % remaining.")

            elif 'open video player' in a or 'open Friday video player' in a or 'show video player' in a or 'show Friday video player' in a:
                speak("Opening Friday video player")
                os.startfile("Multimedia Player\\dist\\Multimediaplayer\\Multimediaplayer.exe")

            # # else:
            #         ## This code below helps to search anything we type not in our list of items, in google and give the information ##
            #         speak("Searching for "+a)
            #         chrome_path = 'C:\\Program Files (x86)\\Google\\Chrome\\Application\\chrome.exe %s'
            #         for url in search(a, tld="co.in", num=1, stop = 1, pause = 2):
            #                 webbrowser.open("https://google.com/search?q=%s" % a)

            #                 print("Links related to your search are given below : ")
            #                 print("")
            #         class Gsearch_python:
            #                 def __init__(self,name_search):
            #                         self.name = name_search
            #                 def Gsearch(self):
            #                         count = 0
            #                         try :
            #                                 from googlesearch import search
            #                         except ImportError:
            #                                 print("No Module named 'google' Found")

            #                         for i in search(query=self.name,tld='co.in',lang='en',num=10,stop=1,pause=2):
            #                                 count += 1
            #                                 print (count)
            #                                 print(i + '\n')
            #         if __name__=='__main__':
            #                 gs = Gsearch_python(a)
            #                 gs.Gsearch()

            elif "code" in a or 'program' in a:
                    speak("Generating Code")
                    await process_code_blocks(a)
            else:
                    resp = await AsyncClient.create_completion("gpt3", a)
                    print(f"🤖: {resp}")
                    speak(resp)

        except sr.UnknownValueError:
            print("👦: Sorry, could not understand audio.")
        except sr.RequestError as e:
            print(f"👦: Could not request results from Google Speech Recognition service; {e}")
        except Exception as e:
            print(f"🤖: {e}")


clear = lambda: os.system('cls')
clear()
run(main())
